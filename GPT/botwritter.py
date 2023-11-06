import os
import openai
import requests
import json
from docx import Document

palavrachave = []

for artigo in palavrachave:
    api_key = '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'
    api_url = 'https://api.openai.com/v1/chat/completions'
    id_modelo = "gpt-3.5-turbo"

    headers = {
    'Authorization': f'Bearer {api_key}',
    'Content-Type': 'application/json'  
    }


    def get_response(messages):
        data = {
            'model': id_modelo,
            'max_tokens': 350,  
            'messages': messages
        }
        response = requests.post(api_url, headers=headers, json=data)
        return response.json()["choices"][0]["message"]["content"]

    # Headline
    messages = [
        {"role": "system", "content": f"Comporte-se como um escritor especialista em SEO. Sua missão é produzir um artigo de 1500 palavras para ranquear para a palavra-chave: {artigo}"},
        {"role": "user", "content": f"Crie uma Headline para ranquear para a palavra chave: {artigo}"}
    ]
    headline = get_response(messages)

    # Lead
    messages.append({"role": "user", "content": f"Escreva um lead sobre o artigo relacionado à palavra-chave: {artigo}."})
    lead = get_response(messages)

    # Subtítulos e parágrafos
    subtitles_and_paragraphs = []
    for i in range(4):  # Exemplo com 4 subtítulos e parágrafos
        messages.append({"role": "user", "content": f"Sugira um subtítulo relacionado à palavra-chave: {artigo}."})
        subtitle = get_response(messages)
        
        messages.append({"role": "user", f"content": "Escreva parágrafos sobre {subtitle}."})
        paragraph = get_response(messages)
        
        subtitles_and_paragraphs.append((subtitle, paragraph))

    # Criando o documento do Word
    doc = Document()
    doc.add_heading(headline, level=1)
    doc.add_paragraph(lead)

    for subtitle, paragraph in subtitles_and_paragraphs:
        doc.add_heading(subtitle, level=2)
        doc.add_paragraph(paragraph)

    doc.save(f"C:\\Users\\lenovo\\Desktop\\GPT\\Artigo - {artigo}.docx")




