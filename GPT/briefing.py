import os
import openai
import requests
import json

api_key = ''
api_url = 'https://api.openai.com/v1/chat/completions'
id_modelo = "gpt-3.5-turbo"

headers = {
    'Authorization': f'Bearer {api_key}',
    'Content-Type': 'application/json'
}

palavrachave = 'Design de Fachada'

main = {
    'model': id_modelo,
    'messages': [

    {"role": "system", "content": f"Comporte-se como um escritor especialista em SEO. Sua missão é produzir artigos voltados para ranquear a {palavrachave}"},
    {"role": "user", "content": f"Faça um Briefing para produzir um artigo otimizado para SEO para ranquear para a {palavrachave}"}

    ]
}
main = json.dumps(main)


requisicao = requests.post(api_url, headers=headers, data=main)

resposta = requisicao.json()
mainmensagem = resposta["choices"][0]["message"]["content"]


from docx import Document

# Crie um novo documento do Word
doc = Document()

# Briefing
doc.add_heading('Briefing', level=1)

# Adicione parágrafos
paragrafo = mainmensagem
doc.add_paragraph(paragrafo)

# Salve o documento
doc.save(f"C:\\Users\\lenovo\\Desktop\\GPT\\Briefing - {palavrachave}.docx")

